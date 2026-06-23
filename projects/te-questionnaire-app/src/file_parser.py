"""
File Parser — reads Excel, CSV, Word, PDF, and plain text files,
returning a normalized structure for question extraction and writing.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any


SUPPORTED_EXTENSIONS = {".xlsx", ".xls", ".csv", ".docx", ".doc", ".pdf", ".txt"}


def parse(path: Path) -> dict[str, Any]:
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

    if ext in (".xlsx", ".xls"):
        return _parse_excel(path)
    if ext == ".csv":
        return _parse_csv(path)
    if ext in (".docx", ".doc"):
        return _parse_word(path)
    if ext == ".pdf":
        return _parse_pdf(path)
    return _parse_text(path)


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------

def _parse_excel(path: Path) -> dict:
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=True)
    sheets: dict[str, list] = {}

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows():
            cells = []
            for cell in row:
                cells.append({
                    "value": cell.value,
                    "row": cell.row,
                    "col": cell.column,
                    "coordinate": cell.coordinate,
                })
            rows.append(cells)
        sheets[sheet_name] = rows

    result: dict = {
        "type": "excel",
        "path": str(path),
        "sheets": sheets,
        "workbook": wb,
    }

    # Check if this workbook matches the RFP evaluation-matrix layout and,
    # if so, embed the column mapping so the extractor and writer can use it.
    rfp = _detect_rfp_matrix(wb)
    if rfp:
        result["rfp_matrix"] = True
        result["rfp_cols"] = rfp

    return result


def _detect_rfp_matrix(wb) -> dict | None:
    """
    Inspect the first worksheet for a ThousandEyes RFP evaluation matrix.

    Signature: row 1 must have 'thousandeyes' (case-insensitive) somewhere in
    columns 3–10, with at least two additional feature-category labels in the
    columns immediately to its right.

    Returns a column-mapping dict on match, or None.
    """
    ws = wb.active
    if ws.max_row < 3 or ws.max_column < 5:
        return None

    # Scan first 3 rows for the header signature
    te_col: int | None = None
    feature_cols: dict[str, int] = {}

    for row_idx in range(1, 4):
        for col_idx in range(1, min(ws.max_column + 1, 20)):
            raw = ws.cell(row_idx, col_idx).value
            if not raw:
                continue
            val = str(raw).strip().lower()
            if val == "thousandeyes":
                te_col = col_idx
                # Collect feature labels in the columns to the right
                for fc in range(col_idx + 1, min(col_idx + 10, ws.max_column + 1)):
                    fv = ws.cell(row_idx, fc).value
                    if fv and str(fv).strip():
                        feature_cols[str(fv).strip().lower()] = fc
                if len(feature_cols) >= 2:
                    break
        if te_col is not None and len(feature_cols) >= 2:
            break

    if te_col is None or len(feature_cols) < 2:
        return None

    # Verify there are data rows with descriptions in col C (col 3)
    # Look for a column A (section) and column B (category), col C (description)
    desc_col = 3   # Col C
    cat_col  = 2   # Col B
    sect_col = 1   # Col A

    data_rows = 0
    for r in range(2, min(ws.max_row + 1, 20)):
        v = ws.cell(r, desc_col).value
        if v and len(str(v).strip()) > 5:
            data_rows += 1
    if data_rows < 3:
        return None

    return {
        "section_col":    sect_col,
        "category_col":   cat_col,
        "description_col": desc_col,
        "support_col":    te_col,
        "feature_cols":   feature_cols,   # e.g. {"synthetics": 6, "device layer": 7, …}
    }


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------

def _parse_csv(path: Path) -> dict:
    import pandas as pd

    df = pd.read_csv(path, dtype=str).fillna("")
    return {
        "type": "csv",
        "path": str(path),
        "dataframe": df,
    }


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

def _parse_word(path: Path) -> dict:
    from docx import Document

    doc = Document(path)

    paragraphs = []
    for idx, para in enumerate(doc.paragraphs):
        paragraphs.append({
            "index": idx,
            "text": para.text.strip(),
            "style": para.style.name,
        })

    tables: list[list[list[dict]]] = []
    for t_idx, table in enumerate(doc.tables):
        table_data: list[list[dict]] = []
        for r_idx, row in enumerate(table.rows):
            row_data: list[dict] = []
            for c_idx, cell in enumerate(row.cells):
                row_data.append({
                    "row": r_idx,
                    "col": c_idx,
                    "text": cell.text.strip(),
                })
            table_data.append(row_data)
        tables.append(table_data)

    result: dict = {
        "type": "word",
        "path": str(path),
        "paragraphs": paragraphs,
        "tables": tables,
        "doc": doc,
    }

    rfp_sections = _detect_word_rfp(paragraphs)
    if rfp_sections:
        result["rfp_word"]     = True
        result["rfp_sections"] = rfp_sections

    return result


# Matches "3.1.1 Solution", "3.1.10 Cost and Licensing", etc.
_WORD_RFP_SECTION_RE = re.compile(r"^\d+\.\d+\.\d+\s+\S")

# Styles that represent bulleted list items
_BULLET_STYLES = {
    "List Paragraph", "List", "List Bullet",
    "DW List Bullet", "List Bullet 2", "List Bullet 3",
}


def _detect_word_rfp(paragraphs: list[dict]) -> list[dict] | None:
    """
    Detect a numbered-section RFP Word document.

    Returns a list of section dicts [{header, bullets:[{idx,text}]}] when
    found, or None if the document doesn't match the pattern.
    """
    sections: list[dict] = []
    current: dict | None = None

    for p in paragraphs:
        text  = p["text"]
        style = p["style"]
        if not text:
            continue

        if _WORD_RFP_SECTION_RE.match(text):
            if current:
                sections.append(current)
            current = {"header": text, "bullets": []}
        elif current is not None and style in _BULLET_STYLES:
            current["bullets"].append({"idx": p["index"], "text": text})

    if current:
        sections.append(current)

    # Require at least 2 numbered sections with bullets
    qualifying = [s for s in sections if s["bullets"]]
    if len(qualifying) < 2:
        return None

    return qualifying


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _parse_pdf(path: Path) -> dict:
    import pdfplumber

    pages: list[dict] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            pages.append({"page": i + 1, "text": text})

    return {
        "type": "pdf",
        "path": str(path),
        "pages": pages,
    }


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------

def _parse_text(path: Path) -> dict:
    content = path.read_text(encoding="utf-8", errors="replace")
    lines = content.splitlines()
    return {
        "type": "text",
        "path": str(path),
        "lines": lines,
        "content": content,
    }
