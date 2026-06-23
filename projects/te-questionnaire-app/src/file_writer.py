"""
File Writer — produces the final output files when the user clicks "Save".

Primary outputs (ALL input types → professional Word document):
  Word (.docx)  — ThousandEyes-branded, table-based Q&A document matching the
                  reference proposal format (#1B4D89 header, #E7F0F7 alt rows)

Secondary / automatic output:
  Excel (.xlsx) — Companion spreadsheet: Question | Answer | Source URL,
                  saved automatically alongside the Word file.

Legacy specialised outputs (also generated when applicable):
  Excel RFP matrix — writes X-marks back into the original evaluation spreadsheet
"""

from __future__ import annotations

import datetime
from pathlib import Path
from typing import Any

from docx.shared import Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Shared colour palette  (matches reference doc)
# ---------------------------------------------------------------------------

_TE_BLUE       = RGBColor(0x1B, 0x4D, 0x89)   # #1B4D89  title / header fill
_TE_LIGHT_BLUE = RGBColor(0xE7, 0xF0, 0xF7)   # #E7F0F7  alternating row tint
_WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
_BLACK         = RGBColor(0x00, 0x00, 0x00)
_MUTED         = RGBColor(0x59, 0x59, 0x59)
_LINK_BLUE     = RGBColor(0x05, 0x63, 0xC1)   # #0563C1  hyperlink colour
_AMBER_TEXT    = RGBColor(0x7A, 0x5C, 0x00)   # #7A5C00  partial / amber text
_GREEN_TEXT    = RGBColor(0x1A, 0x6B, 0x2E)   # #1A6B2E  fully supported accent

# Hex strings used in OOXML shading attributes
_HEX_TE_BLUE       = "1B4D89"
_HEX_TE_LIGHT_BLUE = "E7F0F7"
_HEX_WHITE         = "FFFFFF"
_HEX_AMBER_LIGHT   = "FFF3CD"   # light amber  — partial KPI card background
_HEX_GREY_LIGHT    = "F5F5F5"   # light grey   — not-applicable KPI card background
_HEX_GREEN_LIGHT   = "E6F4EA"   # light green  — fully-supported KPI card background


# ---------------------------------------------------------------------------
# Support-level counting helper
# ---------------------------------------------------------------------------

def _count_support_levels(questions: list[dict]) -> tuple[int, int, int, int]:
    """
    Return (yes, partial, not_applicable, no_rfp_result) counts for a list
    of question dicts.

    Handles both the new three-tier schema (support_level key) and the legacy
    boolean schema (supported key) so cached/older results work correctly.

    no_rfp_result covers pure Q&A questions that have an answer but no
    rfp_result dict at all — they are treated as "yes" for display purposes.
    """
    yes = partial = na = no_rfp = 0
    for q in questions:
        rfp = q.get("rfp_result")
        if rfp is None:
            # Pure Q&A question (no RFP evaluation) — count if answered
            if q.get("answer"):
                no_rfp += 1
            continue

        level = rfp.get("support_level")
        if level == "yes":
            yes += 1
        elif level == "partial":
            partial += 1
        elif level == "not_applicable":
            na += 1
        else:
            # Legacy boolean fallback — rfp_result has "supported" but no "support_level"
            if rfp.get("supported", False):
                yes += 1
            else:
                na += 1

    return yes, partial, na, no_rfp


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def write(results: dict[str, Any], output_path: str) -> None:
    """
    Primary entry point called from main.py on "Save Updated File".

    Always writes a professional ThousandEyes Word document.
    For Excel-RFP inputs the legacy evaluation matrix is ALSO written
    (output_path is already .docx; the matrix path is derived from it).
    """
    parsed    = results["parsed"]
    questions = results["questions"]
    file_type = parsed["type"]

    # Always produce the professional Word document
    _write_professional_word(parsed, questions, output_path)

    # For RFP evaluation-matrix Excel inputs, additionally mark up the
    # original spreadsheet (saved alongside the Word file as _matrix.xlsx)
    if file_type == "excel" and parsed.get("rfp_matrix"):
        matrix_path = str(Path(output_path).with_suffix("")) + "_rfp_matrix.xlsx"
        _write_excel_rfp(parsed, questions, matrix_path)


def write_qa_excel(questions: list[dict[str, Any]], excel_path: str) -> int:
    """
    Write a standalone Q|Answer|Source URL Excel companion file.
    Returns the number of rows written.
    Delegates to qa_reference.save_reference() which already implements
    the branded ThousandEyes Excel format.
    """
    from src.qa_reference import save_reference
    from pathlib import Path as _Path
    return save_reference(questions, _Path(excel_path))


# ---------------------------------------------------------------------------
# Professional Word document  (primary output for ALL input types)
# ---------------------------------------------------------------------------

def _write_professional_word(
    parsed: dict, questions: list[dict], output_path: str
) -> None:
    """
    Produce a ThousandEyes-branded Word document that matches the reference
    proposal format (blue #1B4D89 cover/headers, #E7F0F7 alternating rows,
    section-organised Q&A table).
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()

    # ── Page setup (72 pt = 1 in margins) ─────────────────────────────────
    for sec in doc.sections:
        sec.left_margin   = Pt(72)
        sec.right_margin  = Pt(72)
        sec.top_margin    = Pt(72)
        sec.bottom_margin = Pt(72)

    # ── OOXML helpers ───────────────────────────────────────────────────────

    def _cell_shading(cell, hex_fill: str) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:shd")):
            tcPr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_fill)
        tcPr.append(shd)

    def _cell_borders(cell) -> None:
        """Add thin grey borders to a table cell."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "C5D8EC")
            borders.append(el)
        tcPr.append(borders)

    def _cell_padding(cell, top=60, bottom=60, left=80, right=80) -> None:
        """Set cell internal padding (in twentieths of a point / twips)."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        mar  = OxmlElement("w:tcMar")
        for side, val in (("top", top), ("bottom", bottom),
                          ("left", left), ("right", right)):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"),    str(val))
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        tcPr.append(mar)

    def _add_hyperlink(paragraph, url: str) -> None:
        """Add a URL as coloured, underlined text (no rel required)."""
        r = paragraph.add_run(url)
        r.font.color.rgb = _LINK_BLUE
        r.font.underline = True
        r.font.size      = Pt(9)

    # ── Heading-style helpers ───────────────────────────────────────────────

    def _h1(text: str) -> None:
        try:
            h = doc.add_heading(text, level=1)
        except Exception:
            h = doc.add_paragraph(text)
        h.runs[0].font.color.rgb = _TE_BLUE if h.runs else None

    def _h2(text: str) -> None:
        try:
            h = doc.add_heading(text, level=2)
        except Exception:
            h = doc.add_paragraph(text)
        if h.runs:
            h.runs[0].font.color.rgb = _TE_BLUE

    def _body(text: str, italic: bool = False, color: RGBColor = _BLACK,
              size: float = 11, space_after: float = 6) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size      = Pt(size)
        r.font.color.rgb = color
        r.italic         = italic
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)

    # ── COVER PAGE ──────────────────────────────────────────────────────────

    def _centered(text: str, size: float, bold: bool = False,
                  color: RGBColor = _BLACK, space_after: float = 6) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size      = Pt(size)
        r.bold           = bold
        r.font.color.rgb = color
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)

    # Spacer at top
    doc.add_paragraph()
    doc.add_paragraph()

    _centered("ThousandEyes",               size=24, bold=True,  color=_TE_BLUE, space_after=8)
    _centered("Proposal for",               size=18, bold=True,  color=_BLACK,   space_after=4)

    # Derive a customer name from the source filename
    src_path = parsed.get("path") or parsed.get("filename") or ""
    customer = Path(src_path).stem.replace("_", " ").replace("-", " ").strip() or "Your Organization"
    _centered(customer,                     size=18, bold=True,  color=_TE_BLUE, space_after=8)

    _centered("ENTERPRISE NETWORK INTELLIGENCE", size=16, bold=True, color=_BLACK, space_after=4)

    # RFP # if present in source paragraphs
    rfp_ref = ""
    for p_data in parsed.get("paragraphs", []):
        t = p_data.get("text", "") if isinstance(p_data, dict) else str(p_data)
        if "RFP #" in t or "RFP#" in t:
            rfp_ref = t.strip()
            break
    if rfp_ref:
        _centered(rfp_ref, size=14, bold=True, color=_BLACK, space_after=6)

    doc.add_paragraph()
    date_str = datetime.date.today().strftime("%B %d, %Y")
    _centered(date_str, size=12, bold=False, color=_MUTED, space_after=6)
    doc.add_paragraph()

    _centered("<Account Executive's Name>",  size=11, color=_MUTED)
    _centered("<Account Executive's Email>", size=11, color=_MUTED)
    _centered("<Additional Contact Info>",   size=11, color=_MUTED)

    doc.add_page_break()

    # ── DISCLAIMERS ─────────────────────────────────────────────────────────

    _h2("Disclaimers")
    for disc in [
        "Thank you for the opportunity to submit this proposal. The information "
        "provided is true and accurate to the best of our knowledge as of the "
        "date above and is non-binding except as expressly agreed in writing.",

        "This response was prepared using ThousandEyes documentation sourced "
        "from docs.thousandeyes.com. All answers should be reviewed by a "
        "ThousandEyes Solution Engineer before final submission.",

        "Any information about product roadmap outlines our general direction "
        "and is subject to change at Cisco's sole discretion without notice.",
    ]:
        _body(disc, size=11, space_after=8)

    doc.add_paragraph()

    # ── EXECUTIVE SUMMARY ────────────────────────────────────────────────────

    answered   = [q for q in questions if q.get("answer")]
    unanswered = [q for q in questions if not q.get("answer")]
    total      = len(questions)

    yes_n, partial_n, na_n, no_rfp_n = _count_support_levels(answered)
    has_rfp_data = any(q.get("rfp_result") is not None for q in answered)

    _h1("Executive Summary")

    if has_rfp_data:
        # Three-tier KPI table — one card per support level
        kpi_tbl = doc.add_table(rows=1, cols=3)
        kpi_tbl.style   = "Table Grid"
        kpi_tbl.autofit = False
        kpi_w = Inches(2.17)   # 3 × 2.17" ≈ 6.5" usable width

        kpi_defs = [
            # (label, count, hex_bg, text_color, pct_base)
            ("FULLY SUPPORTED",    yes_n,     _HEX_GREEN_LIGHT, _GREEN_TEXT),
            ("PARTIAL COVERAGE",   partial_n, _HEX_AMBER_LIGHT, _AMBER_TEXT),
            ("NOT APPLICABLE",     na_n,      _HEX_GREY_LIGHT,  _MUTED),
        ]

        for cell, (label, count, hex_bg, txt_color) in zip(
            kpi_tbl.rows[0].cells, kpi_defs
        ):
            cell.width = kpi_w
            _cell_shading(cell, hex_bg)
            _cell_padding(cell, top=120, bottom=120, left=100, right=100)
            _cell_borders(cell)

            pct = f"{count / total * 100:.0f}%" if total else "0%"

            # Label (small, muted, uppercase)
            p_lbl = cell.paragraphs[0]
            p_lbl.alignment = 1   # CENTER
            r_lbl = p_lbl.add_run(label)
            r_lbl.font.size      = Pt(8)
            r_lbl.font.color.rgb = _MUTED
            r_lbl.bold           = True

            # Count (large, bold, coloured)
            p_num = cell.add_paragraph()
            p_num.alignment = 1
            r_num = p_num.add_run(str(count))
            r_num.font.size      = Pt(28)
            r_num.bold           = True
            r_num.font.color.rgb = txt_color

            # Percentage (small, muted)
            p_pct = cell.add_paragraph()
            p_pct.alignment = 1
            r_pct = p_pct.add_run(f"of {total} ({pct})")
            r_pct.font.size      = Pt(9)
            r_pct.font.color.rgb = _MUTED

        doc.add_paragraph()

        # Narrative below the KPI cards
        _body(
            f"ThousandEyes has evaluated {total} requirements in this response. "
            f"{yes_n} are fully supported by native platform capabilities, "
            f"{partial_n} are partially supported (with noted gaps or complementary Cisco tools), "
            f"and {na_n} fall outside the ThousandEyes monitoring scope.",
            size=11, space_after=12,
        )
    else:
        # Pure Q&A mode — no RFP tier data, show simpler summary
        _body(
            f"ThousandEyes has answered {len(answered)} of {total} questions "
            f"in this response document.",
            size=11, space_after=12,
        )

    doc.add_page_break()

    # ── TECHNICAL REQUIREMENTS TABLE ────────────────────────────────────────

    _h1("Technical Requirements Response")

    # Group by section / category
    sections: dict[str, list[dict]] = {}
    for q in answered:
        sec = (q.get("section") or q.get("category") or "Requirements").strip()
        sections.setdefault(sec, []).append(q)

    # Page is 6.5" wide (8.5" − 1" margins × 2)
    # Columns: # | Requirement | Response | Source URL
    COL_W = (Inches(0.35), Inches(2.25), Inches(2.75), Inches(1.15))

    for sec_name, sec_qs in sections.items():
        if len(sections) > 1:
            _h2(sec_name)

        # Section support summary — three-tier breakdown
        sec_yes, sec_partial, sec_na, sec_no_rfp = _count_support_levels(sec_qs)
        sec_total = len(sec_qs)

        # Build a compact "X yes · Y partial · Z n/a" summary line
        parts = []
        if sec_yes:
            parts.append(f"{sec_yes} fully supported")
        if sec_partial:
            parts.append(f"{sec_partial} partial")
        if sec_na:
            parts.append(f"{sec_na} not applicable")
        if sec_no_rfp:
            parts.append(f"{sec_no_rfp} answered")

        summ_text = (
            f"{' · '.join(parts)} of {sec_total} requirements in this section."
            if parts else f"{sec_total} requirements in this section."
        )

        summ_p = doc.add_paragraph()
        summ_r = summ_p.add_run(summ_text)
        summ_r.bold           = True
        summ_r.font.size      = Pt(10)
        # Colour: blue if all fully covered, amber if partial exists, muted if all na
        if sec_partial or sec_na:
            summ_r.font.color.rgb = _AMBER_TEXT if sec_partial else _MUTED
        else:
            summ_r.font.color.rgb = _TE_BLUE
        summ_p.paragraph_format.space_after = Pt(6)

        # ── Q&A Table ─────────────────────────────────────────────────────
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style   = "Table Grid"
        tbl.autofit = False

        # Header row
        hdr_cells = tbl.rows[0].cells
        hdr_labels = ["#", "Requirement / Question",
                      "ThousandEyes Response", "Source"]
        for ci, (cell, label, w) in enumerate(zip(hdr_cells, hdr_labels, COL_W)):
            cell.width = w
            _cell_shading(cell, _HEX_TE_BLUE)
            _cell_padding(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            r = p.add_run(label)
            r.bold           = True
            r.font.color.rgb = _WHITE
            r.font.size      = Pt(10)

        # Data rows
        for row_i, q in enumerate(sec_qs):
            row       = tbl.add_row()
            fill_hex  = _HEX_TE_LIGHT_BLUE if row_i % 2 == 0 else _HEX_WHITE

            q_text     = (q.get("question") or q.get("rfp_requirement") or "").strip()
            answer_txt = (q.get("answer") or "").strip()
            url_txt    = (q.get("source_url") or "").strip()

            for ci, (cell, w) in enumerate(zip(row.cells, COL_W)):
                cell.width = w
                _cell_shading(cell, fill_hex)
                _cell_padding(cell)
                _cell_borders(cell)

            # Col 0: row number
            r0 = row.cells[0].paragraphs[0].add_run(str(row_i + 1))
            r0.font.size      = Pt(9)
            r0.font.color.rgb = _MUTED

            # Col 1: question / requirement  (bold)
            r1 = row.cells[1].paragraphs[0].add_run(q_text)
            r1.font.size = Pt(10)
            r1.bold      = True

            # Col 2: answer
            r2 = row.cells[2].paragraphs[0].add_run(answer_txt)
            r2.font.size = Pt(10)

            # Col 3: source URL
            if url_txt:
                _add_hyperlink(row.cells[3].paragraphs[0], url_txt)

        doc.add_paragraph()

    # ── UNANSWERED / OUT-OF-SCOPE ───────────────────────────────────────────

    if unanswered:
        _h2("Requirements Not Found in ThousandEyes Documentation")
        _body(
            "The following requirements could not be answered using available "
            "ThousandEyes documentation. Please review with your Solution Engineer.",
            size=11, italic=False, space_after=8,
        )
        tbl2 = doc.add_table(rows=1, cols=2)
        tbl2.style = "Table Grid"

        hdr2 = tbl2.rows[0].cells
        for cell, label in zip(hdr2, ["Requirement", "Note"]):
            _cell_shading(cell, _HEX_TE_BLUE)
            _cell_padding(cell, top=80, bottom=80)
            p = cell.paragraphs[0]
            r = p.add_run(label)
            r.bold           = True
            r.font.color.rgb = _WHITE
            r.font.size      = Pt(10)

        for q in unanswered:
            row = tbl2.add_row()
            _cell_shading(row.cells[0], _HEX_TE_LIGHT_BLUE)
            _cell_shading(row.cells[1], _HEX_WHITE)
            _cell_padding(row.cells[0])
            _cell_padding(row.cells[1])

            q_text = (q.get("question") or q.get("rfp_requirement") or "").strip()
            reason = q.get("reason") or "No relevant ThousandEyes documentation found."

            r_q = row.cells[0].paragraphs[0].add_run(q_text)
            r_q.font.size = Pt(10)
            r_q.bold      = True

            r_n = row.cells[1].paragraphs[0].add_run(reason)
            r_n.font.size      = Pt(10)
            r_n.font.color.rgb = _MUTED
            r_n.italic         = True

        doc.add_paragraph()

    # ── CONCLUSION ──────────────────────────────────────────────────────────

    doc.add_page_break()
    _h2("Conclusion")
    _body(
        "ThousandEyes provides comprehensive network intelligence and digital experience "
        "monitoring capabilities that directly address the requirements outlined in this "
        "document. Our unique vantage points across Cloud Agents, Enterprise Agents, and "
        "Endpoint Agents deliver end-to-end visibility that no other vendor can match.",
        size=11, space_after=8,
    )
    _body(
        "We look forward to partnering with your organization to deliver exceptional "
        "network visibility, proactive issue detection, and measurable improvements in "
        "user experience and operational efficiency.",
        size=11, space_after=12,
    )

    p_pow = doc.add_paragraph()
    p_pow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_pow = p_pow.add_run("Powered by ThousandEyes — A Cisco Company")
    r_pow.italic         = True
    r_pow.font.color.rgb = _MUTED
    r_pow.font.size      = Pt(10)

    doc.save(output_path)


# ---------------------------------------------------------------------------
# Excel RFP evaluation matrix  (legacy — marks X in original spreadsheet)
# ---------------------------------------------------------------------------

def _write_excel_rfp(parsed: dict, questions: list[dict], output_path: str) -> None:
    """
    Write RFP evaluation results back into the original spreadsheet cells
    (X marks in the ThousandEyes support and feature-category columns).
    """
    import openpyxl
    from openpyxl.styles import PatternFill, Font, Alignment

    wb = parsed["workbook"]
    ws = wb.active

    green_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
    no_fill    = PatternFill(fill_type=None)
    x_font     = Font(bold=True, color="276221")
    center     = Alignment(horizontal="center", vertical="center")

    managed_cols: set[int] = set()
    for q in questions:
        if not q.get("rfp_requirement"):
            continue
        managed_cols.add(q["support_col"])
        managed_cols.update(q["feature_cols"].values())

    rfp_rows = {q["row"] for q in questions if q.get("rfp_requirement")}
    for row_num in rfp_rows:
        for col_num in managed_cols:
            cell = ws.cell(row_num, col_num)
            cell.value = None
            cell.fill  = no_fill

    for q in questions:
        if not q.get("rfp_requirement"):
            continue
        rfp_result = q.get("rfp_result")
        if rfp_result is None:
            continue

        row_num      = q["row"]
        support_col  = q["support_col"]
        feature_cols = q["feature_cols"]
        supported    = rfp_result.get("supported", False)
        categories   = rfp_result.get("feature_categories", [])

        e_cell = ws.cell(row_num, support_col)
        if supported:
            e_cell.value     = "X"
            e_cell.fill      = green_fill
            e_cell.font      = x_font
            e_cell.alignment = center

        cats_lower = {c.lower().strip() for c in categories}
        for feat_name, col_num in feature_cols.items():
            feat_lower = feat_name.lower().strip()
            matched = any(
                feat_lower == c or feat_lower in c or c in feat_lower
                for c in cats_lower
            )
            if matched:
                f_cell            = ws.cell(row_num, col_num)
                f_cell.value      = "x"
                f_cell.fill       = green_fill
                f_cell.font       = x_font
                f_cell.alignment  = center

    first_cell_val = str(ws.cell(1, 1).value or "").strip()
    if "auto-filled" not in first_cell_val.lower():
        ws.insert_rows(1)
        note = ws.cell(1, 1)
        note.value = (
            "Auto-filled by Customer Questionnaire Crusher — "
            "review before submission"
        )
        note.font = Font(italic=True, color="595959", size=9)

    wb.save(output_path)
